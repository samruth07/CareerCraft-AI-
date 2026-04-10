"""
CareerCraft AI - PDF Resume Parser Tool
Demonstrates: Tools Integration (Syllabus Topic)

Extracts text content from PDF resumes using pdfplumber.
Handles various PDF formats and ensures clean text extraction.
"""

import pdfplumber
import re
import io


def extract_text_from_pdf(file_path: str = None, file_bytes: bytes = None) -> str:
    """
    Extract text content from a PDF resume.
    
    Supports both file path and bytes input (for Chainlit uploads).
    
    Args:
        file_path: Path to the PDF file
        file_bytes: Raw bytes of the PDF file
        
    Returns:
        str: Extracted and cleaned text from the PDF
    """
    text_parts = []
    
    if file_bytes:
        pdf_source = io.BytesIO(file_bytes)
    elif file_path:
        pdf_source = file_path
    else:
        raise ValueError("Either file_path or file_bytes must be provided")
    
    with pdfplumber.open(pdf_source) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    
    full_text = "\n".join(text_parts)
    
    # Clean up the text
    full_text = clean_resume_text(full_text)
    
    return full_text


def clean_resume_text(text: str) -> str:
    """
    Clean extracted resume text by removing artifacts and normalizing whitespace.
    
    Args:
        text: Raw extracted text
        
    Returns:
        str: Cleaned text
    """
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove page numbers
    text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
    
    # Normalize spaces
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Remove leading/trailing whitespace from lines
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    # Remove empty lines at start/end
    text = text.strip()
    
    return text


def extract_text_from_docx(file_path: str = None, file_bytes: bytes = None) -> str:
    """
    Extract text content from a DOCX resume.
    
    Args:
        file_path: Path to the DOCX file
        file_bytes: Raw bytes of the DOCX file
        
    Returns:
        str: Extracted text from the DOCX
    """
    from docx import Document as DocxDocument
    
    if file_bytes:
        doc = DocxDocument(io.BytesIO(file_bytes))
    elif file_path:
        doc = DocxDocument(file_path)
    else:
        raise ValueError("Either file_path or file_bytes must be provided")
    
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text_parts.append(row_text)
    
    return "\n".join(text_parts)


def parse_resume_file(file_path: str = None, file_bytes: bytes = None, 
                       file_name: str = "") -> str:
    """
    Parse a resume file (PDF or DOCX) and extract text.
    
    Args:
        file_path: Path to the resume file
        file_bytes: Raw bytes of the file
        file_name: Original filename (used to detect format)
        
    Returns:
        str: Extracted resume text
    """
    if file_name.lower().endswith('.pdf') or (file_path and file_path.lower().endswith('.pdf')):
        return extract_text_from_pdf(file_path=file_path, file_bytes=file_bytes)
    elif file_name.lower().endswith('.docx') or (file_path and file_path.lower().endswith('.docx')):
        return extract_text_from_docx(file_path=file_path, file_bytes=file_bytes)
    else:
        # Try PDF first, then DOCX
        try:
            return extract_text_from_pdf(file_path=file_path, file_bytes=file_bytes)
        except Exception:
            return extract_text_from_docx(file_path=file_path, file_bytes=file_bytes)
